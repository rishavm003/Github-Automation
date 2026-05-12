import os
from docx import Document
from bot.config import REPORTS_DIR, REVIEW_FILE, ensure_dirs

def generate_reports(date_str, timestamp, artifacts):
    ensure_dirs()
    md_file = os.path.join(REPORTS_DIR, f"Report_{date_str}.md")
    docx_file = os.path.join(REPORTS_DIR, f"Report_{date_str}.docx")
    
    # 1. Generate Markdown Report
    with open(md_file, "w", encoding="utf-8") as f:
        f.write(f"# Daily Execution Report - {date_str}\n\n")
        f.write(f"**Timestamp:** {timestamp}\n")
        f.write(f"**Quote of the Day:** {artifacts.get('quote', {}).get('q', 'N/A')}\n\n")
        
        f.write("## Task Summary Table\n")
        f.write("| Task | Status | Details |\n|---|---|---|\n")
        f.write(f"| Log Commit | Success | Wrote {artifacts.get('log_file', 'No file')} |\n")
        
        readme = artifacts.get('readme_stats')
        if readme and 'error' not in readme:
            f.write(f"| Profile README | Success | {readme.get('total_repos', 0)} Repos, {readme.get('total_stars', 0)} Stars |\n")
        else:
            f.write("| Profile README | Skipped | Token missing or local README not found |\n")
            
        gfi = artifacts.get('good_first_issue')
        if gfi:
            f.write(f"| Good First Issue | Engaged | [{gfi['title']}]({gfi['url']}) |\n")
        else:
            f.write("| Good First Issue | Skipped | No matches or missing token |\n")
            
        claude = artifacts.get('claude_responses', [])
        f.write(f"| Claude Replies | Success | Responded to {len(claude)} issues |\n\n")
        
        if claude:
            f.write("## AI Draft Responses\n")
            for reply in claude:
                f.write(f"- **{reply['repo']} / {reply['issue_title']}**\n  > {reply['draft']}\n\n")
                
        f.write("## Full Artifacts List\n")
        f.write(f"- `data/logs/{date_str}.log`\n")
        f.write(f"- `{md_file}`\n")
        f.write(f"- `{docx_file}`\n")
        
    # 2. Generate Docx Report
    doc = Document()
    doc.add_heading(f"Daily Execution Report - {date_str}", 0)
    doc.add_paragraph(f"Timestamp: {timestamp}")
    doc.add_paragraph(f"Quote of the Day: {artifacts.get('quote', {}).get('q', 'N/A')}")
    
    doc.add_heading("Task Summary", level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = 'Task'
    hdr[1].text = 'Status'
    hdr[2].text = 'Details'
    
    # Add Log Row
    row1 = table.add_row().cells
    row1[0].text = 'Log Commit'
    row1[1].text = 'Success'
    row1[2].text = f"Wrote {artifacts.get('log_file', '')}"
    
    # Save document
    doc.save(docx_file)
    
    return md_file, docx_file

def generate_review_md(drafts):
    if not drafts:
        if os.path.exists(REVIEW_FILE):
            os.remove(REVIEW_FILE)
        return

    with open(REVIEW_FILE, "w", encoding="utf-8") as f:
        f.write("# 👁 Pending AI Drafts for Review\n\n")
        f.write(f"Generated at: {os.popen('date /t').read().strip()} {os.popen('time /t').read().strip()}\n\n")
        f.write("> **How to use:** Review the drafts below. Run `python approve_drafts.py` to post, edit, or reject them.\n\n")
        
        f.write("## Summary\n")
        f.write(f"- Total Pending: **{len(drafts)}**\n\n")
        
        f.write("| # | Repo | Issue Title | Type |\n")
        f.write("|---|------|-------------|------|\n")
        for i, d in enumerate(drafts, 1):
            f.write(f"| {i} | {d['repo']} | [{d['title']}]({d['url']}) | `{d['type']}` |\n")
        
        f.write("\n---\n\n")
        
        for i, d in enumerate(drafts, 1):
            f.write(f"### Draft {i}: {d['repo']}\n")
            f.write(f"**Issue:** [{d['title']}]({d['url']})\n\n")
            f.write("```text\n")
            f.write(d['draft'])
            f.write("\n```\n\n")
            f.write("---\n\n")
